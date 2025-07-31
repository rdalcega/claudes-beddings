#!/usr/bin/env python3
"""
Ingest documents into ChromaDB for semantic search.

WHEN TO RUN THIS:
- Initially: After setting up the embedding system to index all existing content
- Updates: When you add new PDFs, markdown files, or other documents to the project
- Rebuilds: If you want to refresh the entire database (use --rebuild flag)

HOW IT WORKS:
1. Scans your project directory for text files (.md, .pdf, .txt, .rtf, .docx, .html, .json, .xml, .yaml, .rst, .tex, .log, .csv, .tsv)
2. Extracts text content from each file using format-specific extraction methods
3. Breaks large documents into smaller chunks (1000 chars with 100 char overlap)
4. Creates semantic embeddings for each chunk using sentence-transformers
5. Stores everything in ChromaDB with metadata (file path, category, chunk index)
6. Categorizes content automatically (strategy, content, reference, planning, general)
7. Reports any unsupported text files that were detected but not processed

USAGE EXAMPLES:
- python code/embeddings/ingest.py                    # Process current directory
- python code/embeddings/ingest.py --rebuild          # Rebuild database from scratch  
- python code/embeddings/ingest.py -d /path/to/docs   # Process specific directory

The resulting database enables semantic search across all your promotional content,
allowing you to find relevant information even when exact keywords don't match.
"""

import os
import sys
import logging
from pathlib import Path
from typing import List, Dict, Any
import hashlib
import json
import time
import contextlib
import io

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import fitz  # PyMuPDF
from striprtf.striprtf import rtf_to_text
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
import click
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import json
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import yaml
from docx import Document
from docutils.core import publish_parts
from pylatexenc.latex2text import LatexNodes2Text

# Suppress ChromaDB telemetry error messages
logging.getLogger("chromadb.telemetry.product.posthog").setLevel(logging.CRITICAL)

console = Console()


@contextlib.contextmanager
def suppress_system_messages():
    """Context manager to suppress macOS system messages that appear in stderr."""
    # Messages to filter out
    filtered_messages = [
        "Context leak detected",
        "msgtracer returned -1",
        "CoreDuetContext",
        "ContextKit",
    ]
    
    # Capture stderr
    original_stderr = sys.stderr
    captured_stderr = io.StringIO()
    
    try:
        sys.stderr = captured_stderr
        yield
    finally:
        # Restore original stderr
        sys.stderr = original_stderr
        
        # Get captured output and filter it
        captured_output = captured_stderr.getvalue()
        
        if captured_output:
            # Split into lines and filter
            lines = captured_output.splitlines()
            filtered_lines = []
            
            for line in lines:
                # Check if line contains any filtered messages
                should_filter = any(filtered_msg in line for filtered_msg in filtered_messages)
                
                if not should_filter:
                    filtered_lines.append(line)
            
            # Write filtered output to original stderr
            if filtered_lines:
                original_stderr.write('\n'.join(filtered_lines) + '\n')
                original_stderr.flush()

class DocumentIngester:
    def __init__(self, db_path: str = "./code/embeddings/chroma_db", debug: bool = False):
        self.db_path = db_path
        self.debug = debug
        self.cache_file = Path(db_path).parent / ".ingestion_cache.json"
        self.client = chromadb.PersistentClient(
            path=db_path,
            settings=chromadb.Settings(anonymized_telemetry=False)
        )
        self.collection = self.client.get_or_create_collection(
            name="music_promotion_docs",
            metadata={"description": "Music promotion project documents"}
        )
        self.unsupported_files = []  # Track unsupported text files
        self.processed_files = []  # Track all file processing attempts
        self.successful_files = []  # Track successfully processed files
        self.failed_files = []  # Track failed files with reasons
        self.skipped_files = []  # Track files skipped due to no changes
        self.cache = self._load_cache()
    
    def _load_cache(self) -> Dict[str, Any]:
        """Load the ingestion cache from disk."""
        try:
            if self.cache_file.exists():
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    cache = json.load(f)
                    if self.debug:
                        console.print(f"[cyan]DEBUG: Loaded cache with {len(cache.get('files', {}))} entries[/cyan]")
                    return cache
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Failed to load cache: {e}[/cyan]")
        
        # Return empty cache structure
        return {
            "version": "1.0",
            "last_updated": time.time(),
            "files": {}
        }
    
    def _save_cache(self):
        """Save the ingestion cache to disk."""
        try:
            self.cache["last_updated"] = time.time()
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, indent=2)
            if self.debug:
                console.print(f"[cyan]DEBUG: Saved cache with {len(self.cache.get('files', {}))} entries[/cyan]")
        except Exception as e:
            console.print(f"[red]Warning: Failed to save cache: {e}[/red]")
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Calculate a hash of the file's content."""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Failed to hash file {file_path}: {e}[/cyan]")
            return ""
    
    def should_process_file(self, file_path: Path) -> bool:
        """Check if a file should be processed based on changes since last ingestion."""
        try:
            file_key = str(file_path.resolve())
            
            # If file doesn't exist, don't process
            if not file_path.exists():
                return False
            
            # Get current file stats
            stat = file_path.stat()
            current_mtime = stat.st_mtime
            current_size = stat.st_size
            
            # Check if file is in cache
            if file_key not in self.cache["files"]:
                if self.debug:
                    console.print(f"[cyan]DEBUG: File {file_path.name} not in cache - will process[/cyan]")
                return True
            
            cached_info = self.cache["files"][file_key]
            
            # Check if modification time or size changed
            if (cached_info.get("mtime", 0) != current_mtime or 
                cached_info.get("size", 0) != current_size):
                if self.debug:
                    console.print(f"[cyan]DEBUG: File {file_path.name} modified - will process[/cyan]")
                return True
            
            # For extra safety, check content hash for small files (< 1MB)
            if current_size < 1024 * 1024:
                current_hash = self._get_file_hash(file_path)
                cached_hash = cached_info.get("hash", "")
                if current_hash and cached_hash and current_hash != cached_hash:
                    if self.debug:
                        console.print(f"[cyan]DEBUG: File {file_path.name} content changed - will process[/cyan]")
                    return True
            
            if self.debug:
                console.print(f"[cyan]DEBUG: File {file_path.name} unchanged - will skip[/cyan]")
            return False
            
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Error checking file {file_path}: {e} - will process[/cyan]")
            # On error, process the file to be safe
            return True
    
    def _backup_existing_chunks(self, file_path: Path) -> Dict[str, Any]:
        """Backup existing chunks for a file before re-processing (for rollback if needed)."""
        try:
            # Get the relative path that would be stored in metadata
            relative_path = str(file_path.resolve().relative_to(Path.cwd().resolve()))
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Backing up existing chunks for {relative_path}[/cyan]")
            
            # Query for existing documents with this source path
            try:
                existing = self.collection.get(
                    where={"source": relative_path}
                )
                
                if existing['ids']:
                    backup = {
                        'ids': existing['ids'],
                        'documents': existing['documents'],
                        'metadatas': existing['metadatas']
                    }
                    if self.debug:
                        console.print(f"[cyan]DEBUG: Backed up {len(existing['ids'])} existing chunks[/cyan]")
                    return backup
                else:
                    if self.debug:
                        console.print(f"[cyan]DEBUG: No existing chunks found for {relative_path}[/cyan]")
                    return {}
                        
            except Exception as e:
                if self.debug:
                    console.print(f"[cyan]DEBUG: Error querying existing chunks: {e}[/cyan]")
                return {}
                
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Failed to backup existing chunks for {file_path}: {e}[/cyan]")
            return {}
    
    def _remove_existing_chunks(self, file_path: Path, backup_data: Dict[str, Any] = None):
        """Remove existing chunks for a file from ChromaDB."""
        if not backup_data:
            backup_data = self._backup_existing_chunks(file_path)
        
        if not backup_data or not backup_data.get('ids'):
            return
        
        try:
            self.collection.delete(ids=backup_data['ids'])
            if self.debug:
                console.print(f"[cyan]DEBUG: Removed {len(backup_data['ids'])} existing chunks[/cyan]")
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Error deleting existing chunks: {e}[/cyan]")
            raise  # Re-raise to trigger rollback
    
    def _restore_chunks_from_backup(self, backup_data: Dict[str, Any]):
        """Restore chunks from backup data (rollback operation)."""
        if not backup_data or not backup_data.get('ids'):
            return
        
        try:
            self.collection.upsert(
                ids=backup_data['ids'],
                documents=backup_data['documents'],
                metadatas=backup_data['metadatas']
            )
            if self.debug:
                console.print(f"[cyan]DEBUG: Restored {len(backup_data['ids'])} chunks from backup[/cyan]")
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Error restoring chunks from backup: {e}[/cyan]")
            # This is critical - if we can't restore, we've lost data
            console.print(f"[red]CRITICAL: Failed to restore chunks from backup: {e}[/red]")

    def _update_file_cache(self, file_path: Path):
        """Update cache entry for a successfully processed file."""
        try:
            file_key = str(file_path.resolve())
            stat = file_path.stat()
            
            cache_entry = {
                "mtime": stat.st_mtime,
                "size": stat.st_size,
                "processed_at": time.time()
            }
            
            # Add hash for small files
            if stat.st_size < 1024 * 1024:
                cache_entry["hash"] = self._get_file_hash(file_path)
            
            self.cache["files"][file_key] = cache_entry
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Updated cache for {file_path.name}[/cyan]")
                
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Failed to update cache for {file_path}: {e}[/cyan]")
    
    def _remove_from_cache(self, file_path: Path):
        """Remove cache entry for a deleted file."""
        try:
            file_key = str(file_path.resolve())
            
            if file_key in self.cache["files"]:
                del self.cache["files"][file_key]
                self._save_cache()
                
                if self.debug:
                    console.print(f"[cyan]DEBUG: Removed {file_path.name} from cache[/cyan]")
            else:
                if self.debug:
                    console.print(f"[cyan]DEBUG: {file_path.name} was not in cache[/cyan]")
                    
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Failed to remove {file_path} from cache: {e}[/cyan]")
    
    def move_file_in_database(self, old_path: Path, new_path: Path) -> bool:
        """Move a file's database entries to a new path without recreating chunks."""
        try:
            # Get relative paths for both old and new locations
            old_relative = str(old_path.resolve().relative_to(Path.cwd().resolve()))
            new_relative = str(new_path.resolve().relative_to(Path.cwd().resolve()))
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Moving file in database: {old_relative} -> {new_relative}[/cyan]")
            
            # Get existing chunks for the old path
            existing = self.collection.get(where={"source": old_relative})
            
            if not existing['ids']:
                if self.debug:
                    console.print(f"[cyan]DEBUG: No existing chunks found for {old_relative}[/cyan]")
                return False
            
            # Update metadata for all chunks to reflect new path
            updated_metadatas = []
            for metadata in existing['metadatas']:
                updated_metadata = metadata.copy()
                updated_metadata['source'] = new_relative
                updated_metadata['filename'] = new_path.name
                # Update path metadata
                updated_metadata.update(self._extract_path_metadata(new_path))
                updated_metadatas.append(updated_metadata)
            
            # Upsert chunks with updated metadata (same IDs, same content, new metadata)
            self.collection.upsert(
                ids=existing['ids'],
                documents=existing['documents'],
                metadatas=updated_metadatas
            )
            
            # Update cache: remove old entry, add new entry
            old_cache_key = str(old_path.resolve())
            new_cache_key = str(new_path.resolve())
            
            if old_cache_key in self.cache["files"]:
                # Copy cache entry to new key
                self.cache["files"][new_cache_key] = self.cache["files"][old_cache_key].copy()
                # Remove old cache entry
                del self.cache["files"][old_cache_key]
                self._save_cache()
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Successfully moved {len(existing['ids'])} chunks from {old_relative} to {new_relative}[/cyan]")
            
            return True
            
        except Exception as e:
            if self.debug:
                console.print(f"[cyan]DEBUG: Error moving file in database: {e}[/cyan]")
            console.print(f"[red]Error moving {old_path.name} to {new_path.name} in database: {e}[/red]")
            return False
        
    def extract_pdf_text(self, pdf_path: str) -> str:
        """Extract text from PDF file using PyMuPDF."""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text() + "\n"
            doc.close()
            return text.strip()
        except Exception as e:
            console.print(f"[red]Error reading PDF {pdf_path}: {e}[/red]")
            return ""
    
    def extract_rtf_text(self, rtf_path: str) -> str:
        """Extract text from RTF file."""
        try:
            with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as file:
                rtf_content = file.read()
                # Convert RTF to plain text
                text = rtf_to_text(rtf_content, errors='ignore')
                return text.strip()
        except Exception as e:
            console.print(f"[red]Error reading RTF {rtf_path}: {e}[/red]")
            return ""
    
    def extract_docx_text(self, docx_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(docx_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text).strip()
        except Exception as e:
            console.print(f"[red]Error reading DOCX {docx_path}: {e}[/red]")
            return ""
    
    def extract_html_text(self, html_path: str) -> str:
        """Extract text from HTML file."""
        try:
            with open(html_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
                
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text and clean up whitespace
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text.strip()
        except Exception as e:
            console.print(f"[red]Error reading HTML {html_path}: {e}[/red]")
            return ""
    
    def extract_json_text(self, json_path: str) -> str:
        """Extract text values from JSON file."""
        try:
            with open(json_path, 'r', encoding='utf-8', errors='ignore') as file:
                data = json.load(file)
            
            def extract_text_values(obj, texts=None):
                if texts is None:
                    texts = []
                
                if isinstance(obj, dict):
                    for value in obj.values():
                        extract_text_values(value, texts)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_text_values(item, texts)
                elif isinstance(obj, str) and obj.strip():
                    texts.append(obj.strip())
                
                return texts
            
            text_values = extract_text_values(data)
            return '\n'.join(text_values).strip()
        except Exception as e:
            console.print(f"[red]Error reading JSON {json_path}: {e}[/red]")
            return ""
    
    def extract_xml_text(self, xml_path: str) -> str:
        """Extract text from XML file."""
        try:
            tree = ET.parse(xml_path)
            root = tree.getroot()
            
            def extract_element_text(element, texts=None):
                if texts is None:
                    texts = []
                
                if element.text and element.text.strip():
                    texts.append(element.text.strip())
                
                for child in element:
                    extract_element_text(child, texts)
                
                if element.tail and element.tail.strip():
                    texts.append(element.tail.strip())
                
                return texts
            
            text_values = extract_element_text(root)
            return '\n'.join(text_values).strip()
        except Exception as e:
            console.print(f"[red]Error reading XML {xml_path}: {e}[/red]")
            return ""
    
    def extract_yaml_text(self, yaml_path: str) -> str:
        """Extract text values from YAML file."""
        try:
            with open(yaml_path, 'r', encoding='utf-8', errors='ignore') as file:
                data = yaml.safe_load(file)
            
            def extract_text_values(obj, texts=None):
                if texts is None:
                    texts = []
                
                if isinstance(obj, dict):
                    for value in obj.values():
                        extract_text_values(value, texts)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_text_values(item, texts)
                elif isinstance(obj, str) and obj.strip():
                    texts.append(obj.strip())
                
                return texts
            
            text_values = extract_text_values(data)
            return '\n'.join(text_values).strip()
        except Exception as e:
            console.print(f"[red]Error reading YAML {yaml_path}: {e}[/red]")
            return ""
    
    def extract_rst_text(self, rst_path: str) -> str:
        """Extract text from reStructuredText file."""
        try:
            with open(rst_path, 'r', encoding='utf-8', errors='ignore') as file:
                rst_content = file.read()
            
            # Convert RST to plain text
            parts = publish_parts(rst_content, writer_name='html')
            html_content = parts['body']
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)
            
            return text.strip()
        except Exception as e:
            console.print(f"[red]Error reading RST {rst_path}: {e}[/red]")
            return ""
    
    def extract_tex_text(self, tex_path: str) -> str:
        """Extract text from LaTeX file."""
        try:
            with open(tex_path, 'r', encoding='utf-8', errors='ignore') as file:
                tex_content = file.read()
            
            # Convert LaTeX to plain text
            converter = LatexNodes2Text()
            text = converter.latex_to_text(tex_content)
            
            return text.strip()
        except Exception as e:
            console.print(f"[red]Error reading TEX {tex_path}: {e}[/red]")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks."""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to break at sentence boundary
            while end > start + chunk_size // 2 and text[end] not in '.!?\n':
                end -= 1
            
            chunks.append(text[start:end])
            start = end - overlap
            
        return chunks
    
    def process_file(self, file_path: Path, force: bool = False) -> List[Dict[str, Any]]:
        """Process a single file and return document chunks."""
        self.processed_files.append(str(file_path))
        
        # Check if file should be processed (unless force is True)
        if not force and not self.should_process_file(file_path):
            self.skipped_files.append(str(file_path))
            console.print(f"[dim]Skipping {file_path.name} (unchanged)[/dim]")
            return []
        
        if self.debug:
            console.print(f"[cyan]DEBUG: Attempting to process file: {file_path}[/cyan]")
            console.print(f"[cyan]DEBUG: File exists: {file_path.exists()}[/cyan]")
            console.print(f"[cyan]DEBUG: File suffix: {file_path.suffix.lower()}[/cyan]")
            console.print(f"[cyan]DEBUG: File size: {file_path.stat().st_size if file_path.exists() else 'N/A'} bytes[/cyan]")
        
        try:
            console.print(f"[blue]Processing {file_path.name}...[/blue]")
            
            suffix = file_path.suffix.lower()
            
            # Define supported and unsupported formats
            unsupported_text_formats = {'.doc', '.odt', '.pages', '.org', '.adoc', '.asciidoc'}
            
            # Check if this is an unsupported text format
            if suffix in unsupported_text_formats:
                reason = f"Unsupported text format: {suffix}"
                console.print(f"[yellow]{reason}: {file_path.name}[/yellow]")
                if self.debug:
                    console.print(f"[cyan]DEBUG: {reason}[/cyan]")
                self.unsupported_files.append(str(file_path))
                self.failed_files.append({"file": str(file_path), "reason": reason})
                return []
            
            # Extract content based on file type
            if self.debug:
                console.print(f"[cyan]DEBUG: Extracting content using method for {suffix}[/cyan]")
            
            content = ""
            extraction_method = ""
            
            if suffix == '.pdf':
                extraction_method = "PDF extraction (PyMuPDF)"
                content = self.extract_pdf_text(str(file_path))
            elif suffix in ['.md', '.txt']:
                extraction_method = "Plain text reading"
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif suffix == '.rtf':
                extraction_method = "RTF extraction"
                content = self.extract_rtf_text(str(file_path))
            elif suffix == '.docx':
                extraction_method = "DOCX extraction"
                content = self.extract_docx_text(str(file_path))
            elif suffix in ['.html', '.htm']:
                extraction_method = "HTML extraction"
                content = self.extract_html_text(str(file_path))
            elif suffix == '.json':
                extraction_method = "JSON text extraction"
                content = self.extract_json_text(str(file_path))
            elif suffix == '.xml':
                extraction_method = "XML text extraction"
                content = self.extract_xml_text(str(file_path))
            elif suffix in ['.yaml', '.yml']:
                extraction_method = "YAML text extraction"
                content = self.extract_yaml_text(str(file_path))
            elif suffix == '.rst':
                extraction_method = "reStructuredText extraction"
                content = self.extract_rst_text(str(file_path))
            elif suffix == '.tex':
                extraction_method = "LaTeX extraction"
                content = self.extract_tex_text(str(file_path))
            elif suffix in ['.log', '.csv', '.tsv']:
                extraction_method = "Plain text reading"
                # Process as plain text files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            else:
                reason = f"No extraction method for file type: {suffix}"
                if self.debug:
                    console.print(f"[cyan]DEBUG: {reason}[/cyan]")
                self.failed_files.append({"file": str(file_path), "reason": reason})
                return []
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Used extraction method: {extraction_method}[/cyan]")
                console.print(f"[cyan]DEBUG: Extracted content length: {len(content)} characters[/cyan]")
            
            if not content.strip():
                reason = f"Empty content after extraction"
                console.print(f"[yellow]Empty content in {file_path.name}[/yellow]")
                if self.debug:
                    console.print(f"[cyan]DEBUG: {reason}[/cyan]")
                self.failed_files.append({"file": str(file_path), "reason": reason})
                return []
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Content preview: {content[:200]}{'...' if len(content) > 200 else ''}[/cyan]")
            
            chunks = self.chunk_text(content)
            documents = []
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Created {len(chunks)} chunks from content[/cyan]")
            
            for i, chunk in enumerate(chunks):
                doc_id = hashlib.md5(f"{file_path}_{i}".encode()).hexdigest()
                documents.append({
                    'id': doc_id,
                    'content': chunk,
                    'metadata': {
                        'source': str(file_path.resolve().relative_to(Path.cwd().resolve())),
                        'filename': file_path.name,
                        'chunk_index': i,
                        'file_type': file_path.suffix.lower(),
                        'category': self._categorize_file(file_path),
                        **self._extract_path_metadata(file_path)
                    }
                })
            
            if self.debug:
                console.print(f"[cyan]DEBUG: Category assigned: {self._categorize_file(file_path)}[/cyan]")
                console.print(f"[cyan]DEBUG: Source path: {str(file_path.resolve().relative_to(Path.cwd().resolve()))}[/cyan]")
            
            # Implement transactional processing: backup first, then replace
            backup_data = self._backup_existing_chunks(file_path)
            
            try:
                # Remove existing chunks only after we have new ones ready
                if backup_data and backup_data.get('ids'):
                    self._remove_existing_chunks(file_path, backup_data)
                
                console.print(f"[green]✓ Created {len(documents)} chunks from {file_path.name}[/green]")
                self.successful_files.append({"file": str(file_path), "chunks": len(documents), "method": extraction_method})
                
                # Update cache after successful processing
                self._update_file_cache(file_path)
                
                return documents
                
            except Exception as processing_error:
                # Restore backup if processing failed after removal
                if backup_data and backup_data.get('ids'):
                    console.print(f"[yellow]Rolling back changes for {file_path.name} due to processing error[/yellow]")
                    self._restore_chunks_from_backup(backup_data)
                raise processing_error
            
        except Exception as e:
            reason = f"Exception during processing: {str(e)}"
            console.print(f"[red]Error processing {file_path}: {e}[/red]")
            if self.debug:
                console.print(f"[cyan]DEBUG: {reason}[/cyan]")
                import traceback
                console.print(f"[cyan]DEBUG: Traceback: {traceback.format_exc()}[/cyan]")
            self.failed_files.append({"file": str(file_path), "reason": reason})
            return []
    
    def _categorize_file(self, file_path: Path) -> str:
        """Categorize file based on its path."""
        path_str = str(file_path).lower()
        
        if 'strategy' in path_str:
            return 'strategy'
        elif 'lyrics' in path_str or 'analysis' in path_str:
            return 'content'
        elif 'references' in path_str or 'resources' in path_str:
            return 'reference'
        elif 'disorganized' in path_str:
            return 'planning'
        else:
            return 'general'
    
    def _extract_path_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract hierarchical path metadata for efficient pre-query filtering."""
        # Resolve both paths to handle symlinks properly (e.g., /var vs /private/var on macOS)
        resolved_file_path = file_path.resolve()
        resolved_cwd = Path.cwd().resolve()
        relative_path = resolved_file_path.relative_to(resolved_cwd)
        
        # Get path parts (excluding the filename)
        path_parts = list(relative_path.parts[:-1])  # Exclude filename
        
        # Generate ancestor paths for directory matching (as comma-separated string)
        path_ancestors = []
        for i in range(1, len(path_parts) + 1):
            ancestor = '/'.join(path_parts[:i])
            path_ancestors.append(ancestor)
        
        # Create multiple individual metadata fields for each possible directory level
        metadata = {
            'path_depth': len(path_parts),
            'parent_dir': path_parts[-1] if path_parts else '',
            'path_ancestors_str': ','.join(path_ancestors)  # Comma-separated string
        }
        
        # Add individual path level fields for exact matching
        for i, part in enumerate(path_parts):
            metadata[f'path_level_{i}'] = part
        
        return metadata
    
    def check_database_consistency(self) -> Dict[str, Any]:
        """Check for potential duplicate documents in the database."""
        console.print("[blue]Checking database consistency...[/blue]")
        
        try:
            # Get all documents
            all_docs = self.collection.get()
            total_docs = len(all_docs['ids'])
            
            if total_docs == 0:
                console.print("[yellow]Database is empty[/yellow]")
                return {"total_docs": 0, "duplicates": 0, "issues": []}
            
            # Group by source file to check for potential duplicates
            source_groups = {}
            for i, metadata in enumerate(all_docs['metadatas']):
                source = metadata.get('source', 'unknown')
                if source not in source_groups:
                    source_groups[source] = []
                source_groups[source].append({
                    'id': all_docs['ids'][i],
                    'chunk_index': metadata.get('chunk_index', -1),
                    'metadata': metadata
                })
            
            # Check for potential issues
            issues = []
            duplicate_files = []
            orphaned_chunks = []
            
            for source, docs in source_groups.items():
                # Check for duplicate chunk indices within same file
                chunk_indices = [doc['chunk_index'] for doc in docs if doc['chunk_index'] >= 0]
                if len(chunk_indices) != len(set(chunk_indices)):
                    issues.append(f"Duplicate chunk indices in {source}")
                    duplicate_files.append(source)
                
                # Check for orphaned chunks (file no longer exists)
                if source != 'unknown':
                    try:
                        file_path = Path(source)
                        if not file_path.exists():
                            issues.append(f"File no longer exists: {source}")
                            orphaned_chunks.extend([doc['id'] for doc in docs])
                    except Exception:
                        pass
            
            result = {
                "total_docs": total_docs,
                "unique_files": len(source_groups),
                "issues": issues,
                "duplicate_files": duplicate_files,
                "orphaned_chunks": orphaned_chunks
            }
            
            console.print(f"[green]Database consistency check complete:[/green]")
            console.print(f"  Total documents: {total_docs}")
            console.print(f"  Unique source files: {len(source_groups)}")
            if issues:
                console.print(f"  Issues found: {len(issues)}")
                for issue in issues:
                    console.print(f"    - {issue}")
            else:
                console.print("  No issues detected")
            
            return result
            
        except Exception as e:
            console.print(f"[red]Error checking database consistency: {e}[/red]")
            return {"error": str(e)}
    
    def repair_database_issues(self, dry_run: bool = True) -> Dict[str, Any]:
        """Automatically repair common database issues."""
        console.print(f"[blue]{'Analyzing' if dry_run else 'Repairing'} database issues...[/blue]")
        
        try:
            # Check for issues first
            consistency_check = self.check_database_consistency()
            if consistency_check.get("error"):
                return consistency_check
            
            if not consistency_check.get("issues"):
                console.print("[green]No issues found to repair[/green]")
                return {"repaired": 0, "actions": []}
            
            actions = []
            repaired_count = 0
            
            # Remove orphaned chunks (files that no longer exist)
            orphaned_chunks = consistency_check.get("orphaned_chunks", [])
            if orphaned_chunks:
                action_desc = f"Remove {len(orphaned_chunks)} orphaned chunks"
                actions.append(action_desc)
                
                if not dry_run:
                    self.collection.delete(ids=orphaned_chunks)
                    repaired_count += len(orphaned_chunks)
                    console.print(f"[green]✓ Removed {len(orphaned_chunks)} orphaned chunks[/green]")
                else:
                    console.print(f"[yellow]Would remove {len(orphaned_chunks)} orphaned chunks[/yellow]")
            
            # Fix duplicate chunks by reprocessing affected files
            duplicate_files = consistency_check.get("duplicate_files", [])
            for source_file in duplicate_files:
                try:
                    file_path = Path(source_file)
                    if file_path.exists():
                        action_desc = f"Reprocess file with duplicates: {source_file}"
                        actions.append(action_desc)
                        
                        if not dry_run:
                            # Remove existing chunks for this file
                            existing = self.collection.get(where={"source": source_file})
                            if existing['ids']:
                                self.collection.delete(ids=existing['ids'])
                            
                            # Reprocess the file
                            documents = self.process_file(file_path, force=True)
                            if documents:
                                ids = [doc['id'] for doc in documents]
                                doc_contents = [doc['content'] for doc in documents]
                                metadatas = [doc['metadata'] for doc in documents]
                                
                                self.collection.upsert(
                                    ids=ids,
                                    documents=doc_contents,
                                    metadatas=metadatas
                                )
                                repaired_count += 1
                                console.print(f"[green]✓ Reprocessed {source_file}[/green]")
                        else:
                            console.print(f"[yellow]Would reprocess {source_file}[/yellow]")
                            
                except Exception as e:
                    console.print(f"[red]Error repairing {source_file}: {e}[/red]")
            
            # Update cache to reflect repairs
            if not dry_run and repaired_count > 0:
                self._save_cache()
            
            result = {
                "repaired": repaired_count,
                "actions": actions,
                "dry_run": dry_run
            }
            
            if dry_run:
                console.print(f"[blue]Repair analysis complete. {len(actions)} potential actions identified.[/blue]")
                console.print("[dim]Run with dry_run=False to actually repair issues.[/dim]")
            else:
                console.print(f"[green]Repair complete. {repaired_count} issues fixed.[/green]")
            
            return result
            
        except Exception as e:
            console.print(f"[red]Error during database repair: {e}[/red]")
            return {"error": str(e)}
    
    def cleanup_deleted_files(self, project_directory: Path) -> int:
        """Remove database entries for files that no longer exist on disk."""
        console.print("[blue]Checking for deleted files to remove from database...[/blue]")
        
        try:
            # Get all documents from database
            all_docs = self.collection.get()
            if not all_docs['ids']:
                console.print("[dim]No documents in database to check[/dim]")
                return 0
            
            deleted_count = 0
            deleted_files = []
            
            # Group documents by source file
            source_groups = {}
            for i, metadata in enumerate(all_docs['metadatas']):
                source = metadata.get('source', 'unknown')
                if source not in source_groups:
                    source_groups[source] = []
                source_groups[source].append(all_docs['ids'][i])
            
            # Check each unique source file
            for source_path, doc_ids in source_groups.items():
                if source_path == 'unknown':
                    continue
                    
                try:
                    # Construct absolute path from relative source path
                    abs_path = project_directory / source_path
                    
                    if not abs_path.exists():
                        # File has been deleted, remove all its chunks
                        deleted_files.append(source_path)
                        self.collection.delete(ids=doc_ids)
                        deleted_count += len(doc_ids)
                        
                        # Remove from cache
                        cache_key = str(abs_path.resolve())
                        if cache_key in self.cache["files"]:
                            del self.cache["files"][cache_key]
                        
                        if self.debug:
                            console.print(f"[cyan]DEBUG: Removed {len(doc_ids)} chunks for deleted file: {source_path}[/cyan]")
                            
                except Exception as e:
                    if self.debug:
                        console.print(f"[cyan]DEBUG: Error checking file {source_path}: {e}[/cyan]")
                    continue
            
            if deleted_count > 0:
                console.print(f"[yellow]Removed {deleted_count} chunks from {len(deleted_files)} deleted files[/yellow]")
                if self.debug:
                    for file_path in deleted_files:
                        console.print(f"[cyan]DEBUG: Deleted file: {file_path}[/cyan]")
                        
                # Save updated cache
                self._save_cache()
            else:
                console.print("[green]No deleted files found in database[/green]")
            
            return deleted_count
            
        except Exception as e:
            console.print(f"[red]Error during deleted files cleanup: {e}[/red]")
            return 0
    
    def ingest_directory(self, directory: Path, file_patterns: List[str] = None, force: bool = False):
        """Ingest all relevant files from a directory."""
        if file_patterns is None:
            file_patterns = [
                # Original formats
                '**/*.md', '**/*.pdf', '**/*.txt', '**/*.rtf',
                # New supported formats
                '**/*.docx', '**/*.html', '**/*.htm', '**/*.json', '**/*.xml',
                '**/*.yaml', '**/*.yml', '**/*.rst', '**/*.tex',
                '**/*.log', '**/*.csv', '**/*.tsv',
                # Unsupported formats (for detection)
                '**/*.doc', '**/*.odt', '**/*.pages', '**/*.org', '**/*.adoc', '**/*.asciidoc'
            ]
        
        all_files = []
        for pattern in file_patterns:
            all_files.extend(directory.glob(pattern))
        
        # Filter out files we don't want - exclude only code and git directories
        excluded_paths = ['.git/', 'code/']
        
        filtered_files = [
            f for f in all_files 
            if not any(excluded in str(f) for excluded in excluded_paths)
        ]
        
        console.print(f"[green]Found {len(filtered_files)} files to process[/green]")
        
        # Clean up deleted files from database
        self.cleanup_deleted_files(directory)
        
        if self.debug:
            console.print(f"[cyan]DEBUG: File discovery details:[/cyan]")
            console.print(f"[cyan]DEBUG: Total files found by glob patterns: {len(all_files)}[/cyan]")
            console.print(f"[cyan]DEBUG: Files after exclusion filtering: {len(filtered_files)}[/cyan]")
            console.print(f"[cyan]DEBUG: Files to process:[/cyan]")
            for i, f in enumerate(filtered_files):
                console.print(f"[cyan]  {i+1:3d}: {f}[/cyan]")
        
        all_documents = []
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console
        ) as progress:
            task = progress.add_task("Processing files...", total=len(filtered_files))
            
            for file_path in filtered_files:
                progress.update(task, description=f"Processing {file_path.name}")
                documents = self.process_file(file_path, force=force)
                all_documents.extend(documents)
                progress.advance(task)
        
        if all_documents:
            console.print(f"[green]Ingesting {len(all_documents)} document chunks in batches...[/green]")
            
            # Process in smaller batches to avoid memory issues
            batch_size = 50
            total_chunks = len(all_documents)
            
            import time
            
            for i in range(0, total_chunks, batch_size):
                batch = all_documents[i:i + batch_size]
                batch_num = (i // batch_size) + 1
                total_batches = (total_chunks + batch_size - 1) // batch_size
                
                console.print(f"[blue]Processing batch {batch_num}/{total_batches} ({len(batch)} chunks)[/blue]")
                
                try:
                    # Prepare data for ChromaDB with unique ID validation
                    ids = [doc['id'] for doc in batch]
                    documents = [doc['content'] for doc in batch]
                    metadatas = [doc['metadata'] for doc in batch]
                    
                    # Check for duplicate IDs within this batch
                    unique_ids = set(ids)
                    if len(unique_ids) != len(ids):
                        console.print(f"[yellow]Warning: Found duplicate IDs in batch {batch_num}, deduplicating...[/yellow]")
                        # Keep only unique entries (last occurrence wins)
                        seen_ids = set()
                        unique_batch = []
                        for doc in reversed(batch):
                            if doc['id'] not in seen_ids:
                                unique_batch.append(doc)
                                seen_ids.add(doc['id'])
                        # Reverse to maintain original order
                        unique_batch.reverse()
                        
                        # Update batch data
                        ids = [doc['id'] for doc in unique_batch]
                        documents = [doc['content'] for doc in unique_batch]
                        metadatas = [doc['metadata'] for doc in unique_batch]
                        
                        if self.debug:
                            console.print(f"[cyan]DEBUG: Deduplicated batch from {len(batch)} to {len(unique_batch)} documents[/cyan]")
                    
                    # Add delay before processing (except for first batch)
                    if batch_num > 1:
                        console.print(f"[yellow]Waiting 3 seconds before batch {batch_num}...[/yellow]")
                        time.sleep(3)
                    
                    # Upsert to collection with progress indicator
                    console.print(f"[blue]Upserting {len(ids)} documents to ChromaDB...[/blue]")
                    self.collection.upsert(
                        ids=ids,
                        documents=documents,
                        metadatas=metadatas
                    )
                    
                    # Verify the batch was added
                    collection_count = self.collection.count()
                    console.print(f"[green]✓ Batch {batch_num} completed. Total docs in DB: {collection_count}[/green]")
                    
                    # Small delay after successful batch to let ChromaDB settle
                    if batch_num < total_batches:
                        time.sleep(1)
                    
                except Exception as e:
                    console.print(f"[red]Error in batch {batch_num}: {e}[/red]")
                    console.print(f"[yellow]Continuing with next batch...[/yellow]")
                    continue
            
            console.print(f"[green]✓ Successfully ingested {len(all_documents)} document chunks[/green]")
        else:
            console.print("[yellow]No documents to ingest[/yellow]")
        
        # Save cache after processing
        self._save_cache()
        
        # Report skipped files
        if self.skipped_files:
            console.print(f"\n[dim]Skipped {len(self.skipped_files)} unchanged files[/dim]")
            if self.debug:
                for skipped_file in self.skipped_files:
                    console.print(f"  [dim]- {Path(skipped_file).name}[/dim]")
        
        # Report processing results if debug mode is enabled
        if self.debug:
            console.print(f"\n[cyan]DEBUG: Processing Summary:[/cyan]")
            console.print(f"[cyan]DEBUG: Total files attempted: {len(self.processed_files)}[/cyan]")
            console.print(f"[cyan]DEBUG: Successfully processed: {len(self.successful_files)}[/cyan]")
            console.print(f"[cyan]DEBUG: Skipped (unchanged): {len(self.skipped_files)}[/cyan]")
            console.print(f"[cyan]DEBUG: Failed to process: {len(self.failed_files)}[/cyan]")
            console.print(f"[cyan]DEBUG: Unsupported formats: {len(self.unsupported_files)}[/cyan]")
            
            if self.successful_files:
                console.print(f"\n[cyan]DEBUG: Successfully processed files:[/cyan]")
                for success in self.successful_files:
                    console.print(f"[cyan]  ✓ {success['file']} - {success['chunks']} chunks - {success['method']}[/cyan]")
            
            if self.failed_files:
                console.print(f"\n[cyan]DEBUG: Failed files:[/cyan]")
                for failure in self.failed_files:
                    console.print(f"[cyan]  ✗ {failure['file']} - {failure['reason']}[/cyan]")
        
        # Report unsupported files
        if self.unsupported_files:
            console.print(f"\n[yellow]The following {len(self.unsupported_files)} text files could not be ingested:[/yellow]")
            for file_path in self.unsupported_files:
                console.print(f"  [dim]{file_path}[/dim]")
            console.print("[dim]These formats require additional libraries not currently supported.[/dim]")


class DocumentWatcher(FileSystemEventHandler):
    """File system event handler for watching document changes."""
    
    def __init__(self, ingester: DocumentIngester, project_dir: Path, debounce_seconds: float = 5.0, verbose: bool = False):
        super().__init__()
        self.ingester = ingester
        self.project_dir = project_dir
        self.debounce_seconds = debounce_seconds
        self.verbose = verbose
        self.pending_files = {}  # file_path -> last_event_time
        self.pending_deletions = {}  # file_path -> deletion_time (for atomic operation detection)
        self.pending_moves = {}  # old_path -> (new_path, event_time) for move detection
        self.file_hashes = {}  # file_path -> last_known_hash
        self.supported_extensions = {
            '.md', '.pdf', '.txt', '.rtf', '.docx', '.html', '.htm', 
            '.json', '.xml', '.yaml', '.yml', '.rst', '.tex', 
            '.log', '.csv', '.tsv'
        }
        self.excluded_paths = ['.git/', 'code/']
        self.manual_scan_timer = None
        self.scan_interval = 15.0  # Check for missed files every 15 seconds
        self.max_retries = 3  # Maximum retries for failed operations
        self.retry_delay = 1.0  # Initial delay between retries (exponential backoff)
        self.atomic_operation_delay = 2.0  # Wait 2 seconds before processing deletions (to detect atomic operations)
        self.move_detection_window = 10.0  # Time window to correlate delete/create events as moves
        self._initialize_file_hashes()
        self._start_manual_scanning()
    
    def _initialize_file_hashes(self):
        """Initialize file_hashes with current hashes of all supported files to prevent false change detection."""
        if self.ingester.debug:
            console.print("[cyan]DEBUG: Initializing file hashes for watch mode...[/cyan]")
        
        try:
            # Get all supported files in the project directory
            all_files = []
            for ext in self.supported_extensions:
                all_files.extend(self.project_dir.glob(f"**/*{ext}"))
            
            # Filter out excluded paths
            filtered_files = [
                f for f in all_files 
                if not any(excluded in str(f.relative_to(self.project_dir)) for excluded in self.excluded_paths)
            ]
            
            # Initialize hash for each file
            initialized_count = 0
            for file_path in filtered_files:
                if file_path.exists() and file_path.is_file():
                    try:
                        current_hash = self.ingester._get_file_hash(file_path)
                        if current_hash:
                            file_key = str(file_path)
                            self.file_hashes[file_key] = current_hash
                            initialized_count += 1
                            if self.ingester.debug:
                                console.print(f"[cyan]DEBUG: Initialized hash for {file_path.name}[/cyan]")
                    except Exception as e:
                        if self.ingester.debug:
                            console.print(f"[cyan]DEBUG: Failed to hash {file_path.name}: {e}[/cyan]")
            
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Initialized {initialized_count} file hashes for watch mode[/cyan]")
            elif initialized_count > 0:
                console.print(f"[blue]Initialized file tracking for {initialized_count} files[/blue]")
                
        except Exception as e:
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Error during file hash initialization: {e}[/cyan]")
            console.print(f"[yellow]Warning: Could not initialize file tracking: {e}[/yellow]")
    
    def _start_manual_scanning(self):
        """Start periodic manual scanning for missed file changes."""
        import threading
        
        def scan_for_changes():
            try:
                self._manual_scan_files()
            except Exception as e:
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: Manual scan error: {e}[/cyan]")
            finally:
                # Schedule next scan
                if self.manual_scan_timer:
                    self.manual_scan_timer = threading.Timer(self.scan_interval, scan_for_changes)
                    self.manual_scan_timer.start()
        
        self.manual_scan_timer = threading.Timer(self.scan_interval, scan_for_changes)
        self.manual_scan_timer.start()
    
    def _stop_manual_scanning(self):
        """Stop the manual scanning timer and clean up pending operations."""
        if self.manual_scan_timer:
            self.manual_scan_timer.cancel()
            self.manual_scan_timer = None
        
        # Clean up any pending deletions and moves
        if self.pending_deletions:
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Cleaning up {len(self.pending_deletions)} pending deletions[/cyan]")
            self.pending_deletions.clear()
            
        if self.pending_moves:
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Cleaning up {len(self.pending_moves)} pending moves[/cyan]")
            self.pending_moves.clear()
    
    def _manual_scan_files(self):
        """Manually scan for file changes that may have been missed by filesystem events."""
        if self.ingester.debug:
            console.print("[cyan]DEBUG: Running manual file scan for missed changes[/cyan]")
        
        try:
            # Get all supported files in the project directory
            all_files = []
            for ext in self.supported_extensions:
                all_files.extend(self.project_dir.glob(f"**/*{ext}"))
            
            # Filter out excluded paths
            filtered_files = [
                f for f in all_files 
                if not any(excluded in str(f.relative_to(self.project_dir)) for excluded in self.excluded_paths)
            ]
            
            changes_found = 0
            for file_path in filtered_files:
                if self._check_file_changed_by_hash(file_path):
                    changes_found += 1
                    if self.ingester.debug:
                        console.print(f"[cyan]DEBUG: Manual scan detected change: {file_path.name}[/cyan]")
                    self._process_file_change(str(file_path))
            
            if self.ingester.debug and changes_found > 0:
                console.print(f"[cyan]DEBUG: Manual scan found {changes_found} changed files[/cyan]")
                
        except Exception as e:
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Error during manual scan: {e}[/cyan]")
    
    def _check_file_changed_by_hash(self, file_path: Path) -> bool:
        """Check if a file has changed by comparing its hash."""
        if not file_path.exists() or not file_path.is_file():
            return False
        
        try:
            current_hash = self.ingester._get_file_hash(file_path)
            if not current_hash:
                return False
            
            file_key = str(file_path)
            last_hash = self.file_hashes.get(file_key)
            
            if last_hash != current_hash:
                self.file_hashes[file_key] = current_hash
                return True
            
            return False
            
        except Exception as e:
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Hash check error for {file_path.name}: {e}[/cyan]")
            return False
    
    def _schedule_delayed_deletion(self, file_path: str):
        """Schedule a file deletion to be processed after a delay to detect atomic operations."""
        current_time = time.time()
        self.pending_deletions[file_path] = current_time
        
        if self.ingester.debug:
            console.print(f"[cyan]DEBUG: Scheduled delayed deletion for {Path(file_path).name} (for move detection and atomic operations)[/cyan]")
        
        # Schedule processing after delay
        import threading
        timer = threading.Timer(self.atomic_operation_delay, self._process_delayed_deletion, [file_path])
        timer.start()
    
    def _process_delayed_deletion(self, file_path: str):
        """Process a delayed deletion - only proceed if file is still gone and deletion is still pending."""
        current_time = time.time()
        
        # Check if this deletion is still pending (not cancelled by file recreation)
        if file_path not in self.pending_deletions:
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Delayed deletion for {Path(file_path).name} was cancelled (atomic operation detected)[/cyan]")
            return
        
        # Check if enough time has passed since deletion was detected
        deletion_time = self.pending_deletions[file_path]
        if current_time - deletion_time < self.atomic_operation_delay:
            return  # Still within delay period
        
        # Verify file is actually gone
        path_obj = Path(file_path)
        if path_obj.exists():
            # File was recreated - this was an atomic operation
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: File {path_obj.name} was recreated during delay - atomic operation detected[/cyan]")
            self._cancel_pending_deletion(file_path)
            # Treat as modification instead
            self._process_file_change(file_path)
            return
        
        # File is truly deleted - proceed with removal
        try:
            console.print(f"\n[yellow]File permanently deleted: {path_obj.name}[/yellow]")
            
            # Remove chunks from database
            self.ingester._remove_existing_chunks(path_obj)
            
            # Remove from cache
            self.ingester._remove_from_cache(path_obj)
            
            # Remove from file hashes
            if file_path in self.file_hashes:
                del self.file_hashes[file_path]
            
            console.print(f"[green]✓ Removed {path_obj.name} from database and cache[/green]")
            
        except Exception as e:
            console.print(f"[red]Error removing {path_obj.name} from database: {e}[/red]")
        finally:
            # Clean up pending deletion
            self._cancel_pending_deletion(file_path)
    
    def _cancel_pending_deletion(self, file_path: str):
        """Cancel a pending deletion (called when file is recreated during atomic operation)."""
        if file_path in self.pending_deletions:
            del self.pending_deletions[file_path]
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Cancelled pending deletion for {Path(file_path).name}[/cyan]")
    
    def _is_atomic_operation(self, file_path: str) -> bool:
        """Check if a file operation appears to be part of an atomic write operation."""
        # Check if we have a pending deletion for this file (indicates atomic operation)
        if file_path in self.pending_deletions:
            return True
        
        # Check for common atomic operation patterns (temp files)
        path_obj = Path(file_path)
        filename = path_obj.name
        full_path = str(path_obj)
        
        # Common temporary file patterns used in atomic operations
        temp_patterns = ['.tmp', '.temp', '~', '.bak', '.swp', '.swo', '.orig']
        if any(filename.endswith(pattern) for pattern in temp_patterns):
            return True
        
        # Check for embedded temp patterns (like file.tmp.md)
        temp_embedded_patterns = ['.tmp.', '.temp.', '.bak.']
        if any(pattern in filename for pattern in temp_embedded_patterns):
            return True
        
        # Check for macOS/Claude specific temp file patterns
        # Be more specific about macOS temp patterns to avoid false positives
        macos_temp_patterns = [
            '/var/folders/',  # macOS temp directory
            '/private/var/folders/',  # macOS temp directory (full path)
            'TemporaryItems/',  # macOS temporary items
        ]
        for pattern in macos_temp_patterns:
            if pattern in full_path:
                # Additional check: if it's in /var/folders/, look for the /T/ temp marker
                # and ensure it's followed by a UUID-like directory structure
                if '/var/folders/' in pattern:
                    # Look for pattern like /var/folders/.../T/ (not just any /T/)
                    import re
                    temp_pattern = re.compile(r'/var/folders/[^/]+/[^/]+/T/')
                    if temp_pattern.search(full_path):
                        return True
                elif pattern == 'TemporaryItems/' and 'TemporaryItems/' in full_path:
                    return True
                elif '/private/var/folders/' in pattern:
                    # Similar check for private variant
                    import re
                    temp_pattern = re.compile(r'/private/var/folders/[^/]+/[^/]+/T/')
                    if temp_pattern.search(full_path):
                        return True
        
        # Check for hidden temp files (starting with dot)
        if filename.startswith('.') and (filename.endswith('.tmp') or 'tmp' in filename):
            return True
        
        # Check for files with random-looking names (potential temp files)
        # Only check this if we're in a temp directory or filename looks very suspicious
        import re
        temp_dirs = ['/tmp/', '/temp/', 'TemporaryItems', '/var/folders/']
        is_in_temp_dir = any(temp_dir in full_path for temp_dir in temp_dirs)
        
        if is_in_temp_dir:
            # Look for files with long random strings (8+ alphanumeric chars)
            random_pattern = re.compile(r'[a-zA-Z0-9]{8,}')
            if len(filename) > 15 and random_pattern.search(filename):
                # Check if a similar file (without random part) might exist
                for ext in self.supported_extensions:
                    if filename.endswith(ext):
                        # Look for files with same extension in same directory
                        try:
                            similar_files = list(path_obj.parent.glob(f"*{ext}"))
                            if len(similar_files) > 1:  # More than just this file
                                return True
                        except (OSError, PermissionError):
                            # If we can't glob, assume it might be temp
                            return True
                        # If it's a single file with random name in temp dir, probably temp
                        return True
        
        # Check if a similar file (without temp extension) exists
        for pattern in temp_patterns:
            if filename.endswith(pattern):
                original_name = filename[:-len(pattern)]
                original_path = path_obj.parent / original_name
                if original_path.exists():
                    return True
        
        return False
    
    def _detect_file_move(self, created_path: str) -> str:
        """Check if a file creation event is part of a move operation by looking for recent deletions."""
        current_time = time.time()
        created_path_obj = Path(created_path)
        
        # Try to get hash of the created file for better matching
        created_file_hash = None
        if created_path_obj.exists() and created_path_obj.is_file():
            try:
                created_file_hash = self.ingester._get_file_hash(created_path_obj)
            except Exception as e:
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: Could not hash created file {created_path_obj.name}: {e}[/cyan]")
        
        # Look for recent deletions that could be the source of this move
        potential_sources = []
        for deleted_path, deletion_time in list(self.pending_deletions.items()):
            deleted_path_obj = Path(deleted_path)
            
            # Check if deletion happened recently (within move detection window)
            if current_time - deletion_time <= self.move_detection_window:
                match_score = 0
                match_reasons = []
                
                # Primary matching: content hash (strongest indicator)
                if created_file_hash and deleted_path in self.file_hashes:
                    deleted_file_hash = self.file_hashes[deleted_path]
                    if created_file_hash == deleted_file_hash:
                        match_score += 100
                        match_reasons.append("content_hash")
                
                # Secondary matching: exact filename
                if deleted_path_obj.name == created_path_obj.name:
                    match_score += 50
                    match_reasons.append("exact_filename")
                
                # Tertiary matching: same extension
                if deleted_path_obj.suffix == created_path_obj.suffix:
                    match_score += 20
                    match_reasons.append("same_extension")
                
                # Bonus for same file size (if we can get it)
                try:
                    if created_path_obj.exists():
                        created_size = created_path_obj.stat().st_size
                        # Check if we have cached size info
                        cache_key = str(deleted_path_obj.resolve())
                        if cache_key in self.ingester.cache.get("files", {}):
                            cached_info = self.ingester.cache["files"][cache_key]
                            if cached_info.get("size") == created_size:
                                match_score += 30
                                match_reasons.append("same_size")
                except Exception:
                    pass
                
                # Only consider matches with reasonable scores
                if match_score >= 50:  # At least filename match or hash match
                    potential_sources.append((deleted_path, deletion_time, match_score, match_reasons))
                    
                    if self.ingester.debug:
                        console.print(f"[cyan]DEBUG: Potential move candidate: {deleted_path_obj.name} -> {created_path_obj.name} (score: {match_score}, reasons: {match_reasons})[/cyan]")
        
        # Return the best match (highest score, then most recent)
        if potential_sources:
            # Sort by score (highest first), then by time (most recent first)
            potential_sources.sort(key=lambda x: (x[2], x[1]), reverse=True)
            source_path, deletion_time, match_score, match_reasons = potential_sources[0]
            
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Best move match: {Path(source_path).name} -> {created_path_obj.name} (score: {match_score}, reasons: {match_reasons})[/cyan]")
            
            return source_path
        
        return None
    
    def _process_file_move(self, old_path: str, new_path: str):
        """Process a detected file move operation."""
        try:
            old_path_obj = Path(old_path)
            new_path_obj = Path(new_path)
            
            console.print(f"\n[blue]File move detected: {old_path_obj.name} -> {new_path_obj.name}[/blue]")
            
            # Cancel any pending deletions for the source path to prevent conflicts
            if old_path in self.pending_deletions:
                self._cancel_pending_deletion(old_path)
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: Cancelled pending deletion for {old_path_obj.name} due to move[/cyan]")
            
            # Cancel any pending file changes for both paths to avoid conflicts
            if old_path in self.pending_files:
                del self.pending_files[old_path]
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: Cancelled pending changes for old path: {old_path_obj.name}[/cyan]")
            if new_path in self.pending_files:
                del self.pending_files[new_path]
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: Cancelled pending changes for new path: {new_path_obj.name}[/cyan]")
            
            # Verify the destination file exists and is readable
            if not new_path_obj.exists():
                console.print(f"[yellow]Warning: Destination file {new_path_obj.name} does not exist, treating as deletion[/yellow]")
                self._schedule_delayed_deletion(old_path)
                return
            
            if not new_path_obj.is_file():
                console.print(f"[yellow]Warning: Destination {new_path_obj.name} is not a file, treating as deletion[/yellow]")
                self._schedule_delayed_deletion(old_path)
                return
            
            # Try to move the file in the database
            if self.ingester.move_file_in_database(old_path_obj, new_path_obj):
                console.print(f"[green]✓ Moved {old_path_obj.name} to {new_path_obj.name} in database[/green]")
                
                # Update file hash tracking
                if old_path in self.file_hashes:
                    # Transfer the hash to the new path
                    self.file_hashes[new_path] = self.file_hashes[old_path]
                    del self.file_hashes[old_path]
                    if self.ingester.debug:
                        console.print(f"[cyan]DEBUG: Transferred hash tracking from {old_path_obj.name} to {new_path_obj.name}[/cyan]")
                
                # Update current hash for new location to ensure accuracy
                try:
                    current_hash = self.ingester._get_file_hash(new_path_obj)
                    if current_hash:
                        self.file_hashes[new_path] = current_hash
                        if self.ingester.debug:
                            console.print(f"[cyan]DEBUG: Updated hash for {new_path_obj.name} at new location[/cyan]")
                except Exception as hash_error:
                    if self.ingester.debug:
                        console.print(f"[cyan]DEBUG: Could not update hash for {new_path_obj.name}: {hash_error}[/cyan]")
            else:
                # Move failed, fall back to re-processing the file
                console.print(f"[yellow]Database move failed, re-processing {new_path_obj.name}[/yellow]")
                # Clean up old path from database and hash tracking
                if old_path in self.file_hashes:
                    del self.file_hashes[old_path]
                self._process_file_change(new_path)
                
        except Exception as e:
            console.print(f"[red]Error processing file move: {e}[/red]")
            if self.ingester.debug:
                import traceback
                console.print(f"[cyan]DEBUG: Move error traceback: {traceback.format_exc()}[/cyan]")
            # Fall back to re-processing the file
            try:
                # Clean up hash tracking for old path
                if old_path in self.file_hashes:
                    del self.file_hashes[old_path]
            except Exception:
                pass
            self._process_file_change(new_path)
    
    def _should_process_immediately(self, file_path: str) -> bool:
        """Determine if a file should be processed immediately rather than waiting for debouncing."""
        path_obj = Path(file_path)
        filename = path_obj.name
        
        # Skip immediate processing for potential temp files
        if self._is_atomic_operation(file_path):
            return False
        
        # Process immediately for files that appear to be final (not temporary)
        immediate_patterns = [
            # Files created in project directory (not temp dirs)
            lambda p: '/var/folders/' not in str(p) and 'TemporaryItems' not in str(p),
            # Files with normal names (not random temp names)
            lambda p: not any(char in p.name for char in ['~', '.tmp', '.temp', '.bak']),
            # Files that look like typical user files
            lambda p: len(p.name) < 50 and not p.name.startswith('.') and p.suffix in self.supported_extensions,
        ]
        
        # File should be processed immediately if it matches immediate patterns
        if all(pattern(path_obj) for pattern in immediate_patterns):
            # Additional check: file is in the project directory (not a temp location)
            try:
                relative_path = path_obj.relative_to(self.project_dir)
                # File is within project directory and not in excluded paths
                relative_str = str(relative_path)
                if not any(excluded in relative_str for excluded in self.excluded_paths):
                    return True
            except (ValueError, OSError):
                # File is outside project directory or can't be accessed
                pass
        
        return False
    
    def _process_file_change(self, file_path: str):
        """Process a detected file change with event consolidation."""
        current_time = time.time()
        
        # Cancel any pending deletion for this file (atomic operation)
        if file_path in self.pending_deletions:
            self._cancel_pending_deletion(file_path)
            if self.ingester.debug:
                console.print(f"[cyan]DEBUG: Cancelled pending deletion due to file change: {Path(file_path).name}[/cyan]")
        
        # Update pending files timestamp for debouncing
        self.pending_files[file_path] = current_time
        
        # Schedule processing after debounce period
        import threading
        timer = threading.Timer(self.debounce_seconds, self._debounced_process_file, [file_path])
        timer.start()
    
    def _retry_with_backoff(self, operation, *args, max_retries=None, **kwargs):
        """Execute an operation with exponential backoff retry logic."""
        if max_retries is None:
            max_retries = self.max_retries
        
        last_exception = None
        for attempt in range(max_retries + 1):
            try:
                return operation(*args, **kwargs)
            except Exception as e:
                last_exception = e
                if attempt < max_retries:
                    delay = self.retry_delay * (2 ** attempt)  # Exponential backoff
                    if self.ingester.debug:
                        console.print(f"[cyan]DEBUG: Attempt {attempt + 1} failed, retrying in {delay:.1f}s: {e}[/cyan]")
                    time.sleep(delay)
                else:
                    if self.ingester.debug:
                        console.print(f"[cyan]DEBUG: All {max_retries + 1} attempts failed, giving up[/cyan]")
        
        # If we get here, all retries failed
        raise last_exception
    
    def _should_process_file(self, file_path: str) -> bool:
        """Check if a file should be processed based on extension and path."""
        path_obj = Path(file_path)
        
        # Check extension
        if path_obj.suffix.lower() not in self.supported_extensions:
            return False
            
        # Check if in excluded paths
        relative_path = str(path_obj.relative_to(self.project_dir))
        if any(excluded in relative_path for excluded in self.excluded_paths):
            return False
            
        return True
    
    def _debounced_process_file(self, file_path: str):
        """Process a file after debouncing to avoid excessive processing."""
        current_time = time.time()
        
        # Check if enough time has passed since the last event for this file
        if file_path in self.pending_files:
            time_since_last_event = current_time - self.pending_files[file_path]
            if time_since_last_event < self.debounce_seconds:
                return  # Still within debounce period
        
        try:
            path_obj = Path(file_path)
            if not path_obj.exists():
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: File {path_obj.name} no longer exists, skipping processing[/cyan]")
                return
                
            if not path_obj.is_file():
                if self.ingester.debug:
                    console.print(f"[cyan]DEBUG: {path_obj.name} is not a file, skipping processing[/cyan]")
                return
                
            console.print(f"\n[blue]File changed: {path_obj.name}[/blue]")
            
            # Backup existing chunks before any processing
            backup_data = self.ingester._backup_existing_chunks(path_obj)
            
            try:
                # Process the single file to get new chunks
                documents = self.ingester.process_file(path_obj, force=True)
                
                if documents:
                    try:
                        # Add new documents to ChromaDB with retry logic
                        console.print(f"[blue]Adding {len(documents)} chunks to database...[/blue]")
                        
                        ids = [doc['id'] for doc in documents]
                        doc_contents = [doc['content'] for doc in documents]
                        metadatas = [doc['metadata'] for doc in documents]
                        
                        # Use retry logic for database operations
                        def upsert_operation():
                            return self.ingester.collection.upsert(
                                ids=ids,
                                documents=doc_contents,
                                metadatas=metadatas
                            )
                        
                        self._retry_with_backoff(upsert_operation)
                        console.print(f"[green]✓ Updated {path_obj.name} in database[/green]")
                        
                        # Save cache after successful database update with retry
                        def save_cache_operation():
                            return self.ingester._save_cache()
                        
                        try:
                            self._retry_with_backoff(save_cache_operation, max_retries=2)
                        except Exception as cache_error:
                            console.print(f"[yellow]Warning: Failed to save cache for {path_obj.name} after retries: {cache_error}[/yellow]")
                            
                    except Exception as db_error:
                        console.print(f"[red]Database error while updating {path_obj.name} (after retries): {db_error}[/red]")
                        if self.ingester.debug:
                            import traceback
                            console.print(f"[cyan]DEBUG: Database error traceback: {traceback.format_exc()}[/cyan]")
                        
                        # Rollback: restore original chunks if database update failed
                        if backup_data and backup_data.get('ids'):
                            console.print(f"[yellow]Rolling back database changes for {path_obj.name}[/yellow]")
                            try:
                                self._retry_with_backoff(
                                    self.ingester._restore_chunks_from_backup, 
                                    backup_data, 
                                    max_retries=2
                                )
                            except Exception as rollback_error:
                                console.print(f"[red]CRITICAL: Rollback failed for {path_obj.name}: {rollback_error}[/red]")
                        # Don't save cache if database update failed
                        
                else:
                    console.print(f"[dim]No changes needed for {path_obj.name}[/dim]")
                    
            except Exception as processing_error:
                console.print(f"[red]Processing error for {path_obj.name}: {processing_error}[/red]")
                if self.ingester.debug:
                    import traceback
                    console.print(f"[cyan]DEBUG: Processing error traceback: {traceback.format_exc()}[/cyan]")
                
                # Rollback: if processing failed and we removed old chunks, restore them
                if backup_data and backup_data.get('ids'):
                    console.print(f"[yellow]Rolling back changes for {path_obj.name} due to processing failure[/yellow]")
                    self.ingester._restore_chunks_from_backup(backup_data)
                    
        except Exception as e:
            console.print(f"[red]Unexpected error processing {file_path}: {e}[/red]")
            if self.ingester.debug:
                import traceback
                console.print(f"[cyan]DEBUG: Unexpected error traceback: {traceback.format_exc()}[/cyan]")
        finally:
            # Always remove from pending files to prevent memory leaks
            if file_path in self.pending_files:
                del self.pending_files[file_path]
    
    def on_modified(self, event):
        """Handle file modification events."""
        if self.verbose:
            console.print(f"[dim]File system event: MODIFIED {Path(event.src_path).name}[/dim]")
        
        if not event.is_directory and self._should_process_file(event.src_path):
            # Skip temporary files that are part of atomic operations
            if self._is_atomic_operation(event.src_path):
                if self.verbose:
                    console.print(f"[dim]Skipping temp file (atomic operation): {Path(event.src_path).name}[/dim]")
                return
            
            # Update hash for immediate event-based changes
            if self._check_file_changed_by_hash(Path(event.src_path)):
                if self.ingester.debug or self.verbose:
                    console.print(f"[cyan]File system event detected change: {Path(event.src_path).name}[/cyan]")
                self._process_file_change(event.src_path)
            elif self.verbose:
                console.print(f"[dim]No content change detected for {Path(event.src_path).name}[/dim]")
        elif self.verbose:
            console.print(f"[dim]Skipped {Path(event.src_path).name} (not supported or excluded)[/dim]")
    
    def on_created(self, event):
        """Handle file creation events.""" 
        if not event.is_directory and self._should_process_file(event.src_path):
            path_obj = Path(event.src_path)
            
            # Check if this creation cancels a pending deletion (atomic operation)
            if event.src_path in self.pending_deletions:
                console.print(f"\n[blue]Atomic operation detected: {path_obj.name} recreated[/blue]")
                self._cancel_pending_deletion(event.src_path)
                # Process as modification with a small delay to ensure file is fully written
                import threading
                timer = threading.Timer(0.5, self._process_file_change, [event.src_path])
                timer.start()
                return
            
            # Check if this is part of a move operation
            source_path = self._detect_file_move(event.src_path)
            if source_path:
                self._process_file_move(source_path, event.src_path)
                return
            
            # Check for immediate processing patterns (non-atomic operations)
            should_process_immediately = self._should_process_immediately(event.src_path)
            
            if should_process_immediately:
                console.print(f"\n[blue]New file detected (immediate): {path_obj.name}[/blue]")
                # Process immediately with minimal delay
                import threading
                timer = threading.Timer(0.1, self._process_file_change, [event.src_path])
                timer.start()
            else:
                console.print(f"\n[blue]New file detected: {path_obj.name}[/blue]")
                self.on_modified(event)  # Treat creation like modification
    
    def on_moved(self, event):
        """Handle file move events directly."""
        if not event.is_directory:
            src_path = Path(event.src_path)
            dest_path = Path(event.dest_path)
            
            if self.ingester.debug or self.verbose:
                console.print(f"[dim]File system event: MOVED {src_path.name} -> {dest_path.name}[/dim]")
            
            try:
                # Check if both source and destination are files we should process
                should_process_src = self._should_process_file(event.src_path)
                should_process_dest = self._should_process_file(event.dest_path)
                
                if should_process_src and should_process_dest:
                    # Both files are supported - this is a move within our tracked files
                    
                    # Validate that destination file actually exists and is accessible
                    if not dest_path.exists():
                        console.print(f"[yellow]Warning: Move destination {dest_path.name} does not exist, treating as deletion[/yellow]")
                        self._schedule_delayed_deletion(event.src_path)
                        return
                    
                    if not dest_path.is_file():
                        console.print(f"[yellow]Warning: Move destination {dest_path.name} is not a file, treating as deletion[/yellow]")
                        self._schedule_delayed_deletion(event.src_path)
                        return
                    
                    # Check file accessibility
                    try:
                        # Quick read test to ensure file is accessible
                        with open(dest_path, 'rb') as f:
                            f.read(1)  # Read one byte to test accessibility
                    except Exception as access_error:
                        console.print(f"[yellow]Warning: Cannot access moved file {dest_path.name}: {access_error}[/yellow]")
                        # Delay processing to allow filesystem to settle
                        import threading
                        timer = threading.Timer(2.0, self._process_file_change, [event.dest_path])
                        timer.start()
                        return
                    
                    console.print(f"\n[blue]Direct move detected: {src_path.name} -> {dest_path.name}[/blue]")
                    
                    # Cancel any pending deletions for source path
                    if event.src_path in self.pending_deletions:
                        self._cancel_pending_deletion(event.src_path)
                    
                    # Process the move
                    self._process_file_move(event.src_path, event.dest_path)
                    
                elif should_process_src and not should_process_dest:
                    # File moved out of tracked area - treat as deletion
                    console.print(f"\n[yellow]File moved out of tracked area: {src_path.name}[/yellow]")
                    self._schedule_delayed_deletion(event.src_path)
                    
                elif not should_process_src and should_process_dest:
                    # File moved into tracked area - treat as creation
                    console.print(f"\n[blue]File moved into tracked area: {dest_path.name}[/blue]")
                    
                    # Validate the new file before processing
                    if dest_path.exists() and dest_path.is_file():
                        # Small delay to ensure file is fully written
                        import threading
                        timer = threading.Timer(0.5, self._process_file_change, [event.dest_path])
                        timer.start()
                    else:
                        console.print(f"[yellow]Warning: Moved file {dest_path.name} is not accessible, skipping[/yellow]")
                    
                elif self.verbose:
                    # Neither should be processed
                    console.print(f"[dim]Ignored move: {src_path.name} -> {dest_path.name} (not in tracked area)[/dim]")
                    
            except Exception as e:
                console.print(f"[red]Error handling move event {src_path.name} -> {dest_path.name}: {e}[/red]")
                if self.ingester.debug:
                    import traceback
                    console.print(f"[cyan]DEBUG: Move event error traceback: {traceback.format_exc()}[/cyan]")
                
                # Fallback: if source was tracked, treat as deletion + creation
                if self._should_process_file(event.src_path):
                    console.print(f"[yellow]Fallback: treating as deletion of {src_path.name}[/yellow]")
                    self._schedule_delayed_deletion(event.src_path)
                
                if self._should_process_file(event.dest_path):
                    console.print(f"[yellow]Fallback: treating as creation of {dest_path.name}[/yellow]")
                    import threading
                    timer = threading.Timer(1.0, self._process_file_change, [event.dest_path])
                    timer.start()
    
    def on_deleted(self, event):
        """Handle file deletion events with delayed processing to detect atomic operations."""
        if not event.is_directory and self._should_process_file(event.src_path):
            path_obj = Path(event.src_path)
            
            # Don't process deletion immediately - could be part of atomic operation
            if self.ingester.debug or self.verbose:
                console.print(f"\n[yellow]File deletion detected: {path_obj.name} (processing delayed)[/yellow]")
            
            # Schedule delayed deletion processing
            self._schedule_delayed_deletion(event.src_path)


@click.command()
@click.option('--directory', '-d', default='.', help='Directory to ingest documents from')
@click.option('--rebuild', '-r', is_flag=True, help='Rebuild the database from scratch')
@click.option('--force', '-f', is_flag=True, help='Force re-processing of all files, ignoring cache')
@click.option('--watch', '-w', is_flag=True, help='Watch for file changes and automatically re-ingest')
@click.option('--debug', is_flag=True, help='Enable debug logging with detailed file processing information')
@click.option('--check', '-c', is_flag=True, help='Check database consistency and exit')
@click.option('--repair', is_flag=True, help='Check for and repair database issues')
@click.option('--dry-run', is_flag=True, help='Show what repair would do without making changes (use with --repair)')
@click.option('--verbose', '-v', is_flag=True, help='Enable verbose output showing detailed file system events')
def main(directory: str, rebuild: bool, force: bool, watch: bool, debug: bool, check: bool, repair: bool, dry_run: bool, verbose: bool):
    """Ingest documents into the embeddings database."""
    
    project_dir = Path(directory).resolve()
    console.print(f"[blue]Ingesting documents from: {project_dir}[/blue]")
    
    # Use code/embeddings directory for database
    code_embeddings_dir = project_dir / "code" / "embeddings"
    code_embeddings_dir.mkdir(exist_ok=True)
    
    db_path = code_embeddings_dir / "chroma_db"
    
    if rebuild and db_path.exists():
        console.print("[yellow]Rebuilding database...[/yellow]")
        import shutil
        shutil.rmtree(db_path)
        # Also clear the cache file when rebuilding
        cache_file = code_embeddings_dir / ".ingestion_cache.json"
        if cache_file.exists():
            cache_file.unlink()
            console.print("[yellow]Cleared ingestion cache[/yellow]")
    
    ingester = DocumentIngester(str(db_path), debug=debug)
    
    # If check flag is set, just run consistency check and exit
    if check:
        ingester.check_database_consistency()
        return
    
    # If repair flag is set, run repair and exit
    if repair:
        ingester.repair_database_issues(dry_run=dry_run)
        return
    
    # Always do an initial ingestion
    ingester.ingest_directory(project_dir, force=force or rebuild)
    console.print("[green]✓ Initial ingestion complete![/green]")
    
    # Start watching if requested
    if watch:
        console.print("\n[blue]Starting file watcher...[/blue]")
        console.print("[dim]Press Ctrl+C to stop watching[/dim]")
        
        try:
            # Create event handler and observer
            event_handler = DocumentWatcher(ingester, project_dir, verbose=verbose)
            observer = Observer()
            observer.schedule(event_handler, str(project_dir), recursive=True)
            
            # Start watching
            observer.start()
            console.print(f"[green]✓ Watching {project_dir} for changes...[/green]")
            
            # Keep the script running with message filtering
            try:
                with suppress_system_messages():
                    while True:
                        time.sleep(1)
            except KeyboardInterrupt:
                console.print("\n[yellow]Stopping file watcher...[/yellow]")
                observer.stop()
                
        except Exception as e:
            console.print(f"[red]Error starting file watcher: {e}[/red]")
            return
        finally:
            if 'observer' in locals():
                observer.join()
            if 'event_handler' in locals():
                event_handler._stop_manual_scanning()
                console.print("[green]✓ File watcher stopped[/green]")
    else:
        console.print("[green]✓ Ingestion complete![/green]")

if __name__ == "__main__":
    main()